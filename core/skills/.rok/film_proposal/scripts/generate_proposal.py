#!/usr/bin/env python3
"""
ROKCT STUDIOS — Film Proposal Generator

Reads the production bible from film/{project}/ in the consumer repository
and writes a broadcast-ready eTV submission .docx

IMPORTANT: run this script from the consumer repo root. The repo root is
resolved from the current working directory (NOT from the script location —
provisioned skill scripts execute from variable locations).

Usage (from consumer repo root, after .rokct/initiate.py has provisioned skills):
  python3 .rokct/skills/.rok/film_proposal/scripts/generate_proposal.py [project] [out.docx]

Defaults:
  project  -> venda_nga_december
  out.docx -> film/{project}/proposal.docx

What it reads:
  film/{project}/00_index.md
  film/{project}/metarules/world_rules.md
  film/{project}/characters/*.md
  film/{project}/scenes/all_scenes.md
  film/{project}/themes/all_themes.md
  film/{project}/bubbles/*.md

Update any bible file -> rerun -> fresh proposal.

Requires: python-docx  (pip install python-docx)
"""

import os
import re
import sys
import zipfile

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Twips

# ─── PATHS ───────────────────────────────────────────────────────────────────
project = sys.argv[1] if len(sys.argv) > 1 else 'venda_nga_december'
output_arg = sys.argv[2] if len(sys.argv) > 2 else None
# Repo root is the current working directory — run from the consumer repo root.
REPO_ROOT = os.getcwd()
BIBLE = os.path.join(REPO_ROOT, 'film', project)
OUT = os.path.abspath(output_arg) if output_arg else os.path.join(BIBLE, 'proposal.docx')

if not os.path.isdir(BIBLE):
    sys.stderr.write('\nERROR: Cannot find bible folder:\n  %s\n\n' % BIBLE)
    sys.stderr.write('Run this script from the consumer repo root '
                     '(the directory containing film/).\n')
    sys.exit(1)
print('\nReading bible: %s' % BIBLE)


# ─── READ ────────────────────────────────────────────────────────────────────
def read(rel):
    f = os.path.join(BIBLE, rel)
    if os.path.isfile(f):
        with open(f, encoding='utf-8') as fh:
            return fh.read()
    return ''


def read_chars():
    d = os.path.join(BIBLE, 'characters')
    if not os.path.isdir(d):
        return {}
    out = {}
    for f in sorted(os.listdir(d)):
        if f.endswith('.md'):
            with open(os.path.join(d, f), encoding='utf-8') as fh:
                out[f[:-3]] = fh.read()
    return out


IDX = read('00_index.md')
WORLD = read('metarules/world_rules.md')
SCENES = read('scenes/all_scenes.md')
CHARS = read_chars()


# ─── PARSE ───────────────────────────────────────────────────────────────────
def section(text, heading):
    m = re.search(r'##+ %s[\s\S]*?(?=\n##+ |$)' % re.escape(heading), text, re.I)
    return re.sub(r'^##+ [^\n]+\n', '', m.group(0)).strip() if m else ''


def strip_md(t):
    t = re.sub(r'\*\*(.+?)\*\*', r'\1', t)
    t = re.sub(r'\*(.+?)\*', r'\1', t)
    t = re.sub(r'^#+\s+', '', t, flags=re.M)
    t = re.sub(r'^[-*]\s+', '', t, flags=re.M)
    t = re.sub(r'`([^`]+)`', r'\1', t)
    t = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', t)
    return t.strip()


def paras(t):
    return [p for p in
            (p.replace('\n', ' ').strip() for p in re.split(r'\n\n+', strip_md(t)))
            if p]


def get_logline():
    ps = paras(section(IDX, 'The Logline'))
    return ps[0] if ps else ('A man from Thohoyandou tattooed his dream on his chest '
                             'before the world knew his name. This documentary is as '
                             'permanent as that tattoo.')


def get_company():
    m = re.search(r'ROKCT INTELLIGENCE[^\n]*', IDX, re.I)
    return m.group(0).strip() if m else 'ROKCT INTELLIGENCE (PTY) LTD t/a ROKCT STUDIOS'


def get_cq():
    ps = paras(section(WORLD, 'The Central Question'))
    return ps[0] if ps else ('What does a man from Thohoyandou have to carry, sacrifice, '
                             'and survive to make the world hear him — and what does he '
                             'hold onto when everything else is negotiable?')


def char_data(slug):
    t = CHARS.get(slug)
    if not t:
        return None
    m = re.search(r'^# (.+)$', t, re.M)
    name = m.group(1).strip() if m else 'Unknown'
    role = ' '.join(paras(section(t, 'Role'))[:2])
    m = re.search(r'\*\*Fee:\*\*\s*(R[\d,]+)', t, re.I)
    fee = m.group(1) if m else '—'
    m = re.search(r'\*\*Identity:\*\*\s*([^\n]+)', t, re.I)
    ident = m.group(1).strip() if m else 'Negotiated'
    arch = [re.sub(r'^[-*]\s*', '', strip_md(l))
            for l in section(t, 'Archive Bubbles — Verified Facts').split('\n')
            if re.match(r'^[-*]', l.strip())][:4]
    return {'name': name, 'role': role, 'fee': fee, 'id': ident, 'arch': arch}


def scene_block(n):
    # Accept zero-padded headings ("# Scene 01") as well as "# Scene 1";
    # (?!\d) stops "Scene 1" from also matching "Scene 10".
    m = re.search(r'# Scene 0*%d(?!\d)[^\n]*\n([\s\S]*?)(?=\n# Scene |$)' % n, SCENES)
    if not m:
        return None
    blk = m.group(0)
    tm = re.search(r'^# Scene \d+ — (.+)$', blk, re.M)
    title = tm.group(1).strip() if tm else ''
    rm = re.search(r'## Runtime\n([^\n]+)', blk)
    runtime = rm.group(1).strip() if rm else ''
    desc = ' '.join(paras(section(blk, 'What Happens'))[:2])
    return {'title': title, 'runtime': runtime, 'desc': desc}


# ─── DESIGN ──────────────────────────────────────────────────────────────────
DARK, WHITE, ACCENT = '1A1A1A', 'FFFFFF', 'C0392B'
MID, LIGHT, RULEC = '555555', 'F5F5F5', 'CCCCCC'

BULLET_NUM_ID = 100  # numbering definition added below (en-dash bullets)

# Successor tags used to insert pPr children at their schema-valid position.
_PPR_SUCCESSORS = ('w:spacing', 'w:ind', 'w:jc', 'w:rPr', 'w:sectPr')

doc = Document()

# Default document font: Arial 10pt (the JS docx used half-point size 20)
_normal = doc.styles['Normal']
_normal.font.name = 'Arial'
_normal.font.size = Pt(10)

# Heading styles matching the JS paragraphStyles
for _sid, _sz, _col, _before, _after in (
        ('Heading 1', 14, DARK, 18, 6),
        ('Heading 2', 11, ACCENT, 12, 4)):
    _st = doc.styles[_sid]
    _st.font.name = 'Arial'
    _st.font.size = Pt(_sz)
    _st.font.bold = True
    _st.font.color.rgb = RGBColor.from_string(_col)
    _st.paragraph_format.space_before = Pt(_before)
    _st.paragraph_format.space_after = Pt(_after)

# A4 page, JS margins (twips)
_sec = doc.sections[0]
_sec.page_width = Twips(11906)
_sec.page_height = Twips(16838)
_sec.top_margin = Twips(1440)
_sec.right_margin = Twips(1260)
_sec.bottom_margin = Twips(1440)
_sec.left_margin = Twips(1260)

# The python-docx default template ships a <w:zoom> without the required
# w:percent attribute; repair it so the document validates cleanly.
_zoom = doc.settings.element.find(qn('w:zoom'))
if _zoom is not None and _zoom.get(qn('w:percent')) is None:
    _zoom.set(qn('w:percent'), '100')


def _add_bullet_numbering():
    """Register an en-dash bullet list (mirrors the JS 'bullets' numbering config)."""
    numbering = doc.part.numbering_part.element
    abstract = OxmlElement('w:abstractNum')
    abstract.set(qn('w:abstractNumId'), str(BULLET_NUM_ID))
    lvl = OxmlElement('w:lvl')
    lvl.set(qn('w:ilvl'), '0')
    for tag, attrs in (('w:start', {'w:val': '1'}),
                       ('w:numFmt', {'w:val': 'bullet'}),
                       ('w:lvlText', {'w:val': '–'}),
                       ('w:lvlJc', {'w:val': 'left'})):
        el = OxmlElement(tag)
        for k, v in attrs.items():
            el.set(qn(k), v)
        lvl.append(el)
    ppr = OxmlElement('w:pPr')
    ind = OxmlElement('w:ind')
    ind.set(qn('w:left'), '720')
    ind.set(qn('w:hanging'), '360')
    ppr.append(ind)
    lvl.append(ppr)
    abstract.append(lvl)
    numbering.insert(0, abstract)
    num = OxmlElement('w:num')
    num.set(qn('w:numId'), str(BULLET_NUM_ID))
    ref = OxmlElement('w:abstractNumId')
    ref.set(qn('w:val'), str(BULLET_NUM_ID))
    num.append(ref)
    numbering.append(num)


_add_bullet_numbering()


def _run(p, text, size=10, color=DARK, bold=False, italic=False, font='Arial'):
    r = p.add_run(text)
    r.font.name = font
    r.font.size = Pt(size)
    r.font.color.rgb = RGBColor.from_string(color)
    r.bold = bold
    r.italic = italic
    return r


def sp(pt=120):
    """Empty spacer paragraph; pt is in twentieths of a point (as in the JS)."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(pt / 20)
    return p


def hr(color=RULEC, sz=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(8)
    ppr = p._p.get_or_add_pPr()
    pbdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), str(sz))
    bottom.set(qn('w:color'), color)
    pbdr.append(bottom)
    # pBdr must precede spacing in the pPr child sequence
    ppr.insert_element_before(pbdr, *_PPR_SUCCESSORS)
    return p


def h1(t):
    p = doc.add_paragraph(style='Heading 1')
    _run(p, t, size=14, color=DARK, bold=True)
    return p


def h2(t):
    p = doc.add_paragraph(style='Heading 2')
    _run(p, t, size=11, color=ACCENT, bold=True)
    return p


def h3(t):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    _run(p, t, size=10, color=DARK, bold=True)
    return p


def body(t, color=DARK, bold=False, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    _run(p, t, size=10, color=color, bold=bold, italic=italic)
    return p


def bul(t):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    ppr = p._p.get_or_add_pPr()
    numpr = OxmlElement('w:numPr')
    ilvl = OxmlElement('w:ilvl')
    ilvl.set(qn('w:val'), '0')
    numid = OxmlElement('w:numId')
    numid.set(qn('w:val'), str(BULLET_NUM_ID))
    numpr.append(ilvl)
    numpr.append(numid)
    # numPr must precede pBdr/spacing in the pPr child sequence
    ppr.insert_element_before(numpr, 'w:pBdr', *_PPR_SUCCESSORS)
    _run(p, t, size=10, color=DARK)
    return p


def centred(t, sz=20, color=DARK, bold=False, italic=False, before=0, after=80):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(before / 20)
    p.paragraph_format.space_after = Pt(after / 20)
    _run(p, t, size=sz / 2, color=color, bold=bold, italic=italic)
    return p


def page_break():
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)
    return p


def _tc_set(cell, fill=WHITE, no_border=False, pad_left=None, pad_right=None):
    tcpr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement('w:tcBorders')
    for side in ('top', 'left', 'bottom', 'right'):
        el = OxmlElement('w:' + side)
        if no_border:
            el.set(qn('w:val'), 'none')
            el.set(qn('w:sz'), '0')
            el.set(qn('w:color'), WHITE)
        else:
            el.set(qn('w:val'), 'single')
            el.set(qn('w:sz'), '1')
            el.set(qn('w:color'), RULEC)
        borders.append(el)
    tcpr.append(borders)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:fill'), fill)
    tcpr.append(shd)
    mar = OxmlElement('w:tcMar')
    for side, w in (('top', 60),
                    ('left', pad_left if pad_left is not None else 120),
                    ('bottom', 60),
                    ('right', pad_right if pad_right is not None else 120)):
        el = OxmlElement('w:' + side)
        el.set(qn('w:w'), str(w))
        el.set(qn('w:type'), 'dxa')
        mar.append(el)
    tcpr.append(mar)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def cl(cell, text, width, fill=WHITE, color=DARK, bold=False, italic=False,
       align=None, size=18, no_border=False, pad_left=None, pad_right=None):
    cell.width = Twips(width)
    _tc_set(cell, fill=fill, no_border=no_border,
            pad_left=pad_left, pad_right=pad_right)
    p = cell.paragraphs[0]
    if align is not None:
        p.alignment = align
    _run(p, text, size=size / 2, color=color, bold=bold, italic=italic)


def table(col_widths, n_rows):
    t = doc.add_table(rows=n_rows, cols=len(col_widths))
    t.autofit = False  # sets <w:tblLayout w:type="fixed"/> in schema position
    tblpr = t._tbl.tblPr
    tblw = tblpr.find(qn('w:tblW'))
    if tblw is None:
        tblw = OxmlElement('w:tblW')
        tblpr.insert_element_before(
            tblw, 'w:tblJc', 'w:tblCellSpacing', 'w:tblInd', 'w:tblBorders',
            'w:shd', 'w:tblLayout', 'w:tblCellMar', 'w:tblLook')
    tblw.set(qn('w:w'), str(sum(col_widths)))
    tblw.set(qn('w:type'), 'dxa')
    for i, w in enumerate(col_widths):
        t.columns[i].width = Twips(w)
    return t


def h_row(row, cols):
    for cell, (text, width) in zip(row.cells, cols):
        cl(cell, text, width, fill=DARK, color=WHITE, bold=True)


def d_row(row, cols, shade):
    for cell, spec in zip(row.cells, cols):
        text, width = spec[0], spec[1]
        o = spec[2] if len(spec) > 2 else {}
        cl(cell, text, width, fill=(LIGHT if shade else WHITE), **o)


# ─── BUDGET DATA ─────────────────────────────────────────────────────────────
BUDGET = [
    ['Producer / Director fee', 'Principal creative and production leadership', 'R45,000'],
    ['Camera equipment rental', 'Multi-camera locked setup, lenses, supports — 6 days', 'R35,000'],
    ['Sound recordist + equipment', '6 shoot days', 'R18,000'],
    ['Lighting equipment rental', 'Production studio lighting setup', 'R12,000'],
    ['Editor fee', 'Assembly through picture lock incl. bubble layer composition', 'R32,000'],
    ['Colour grade', '3-grade system — studio, b-roll, reconstruction', 'R12,000'],
    ['Sound mix', 'Final broadcast mix', 'R10,000'],
    ['Music licensing / composer', 'VenRap track clearances or original score', 'R14,000'],
    ['Production studio rental + dress', 'Thohoyandou studio hire and preparation', 'R8,000'],
    ['Transport + fuel', 'Local Thohoyandou transport, 6 shoot days', 'R6,000'],
    ['Subtitling', 'English subtitles — professional service', 'R8,000'],
    ['Participant fees', '6 principal contributors', 'R32,000'],
    ['Dramatic reconstruction', 'Local performers, 1 shoot day', 'R8,000'],
    ['Archive / social media research', 'Ricky Rick archive, community bubble collection', 'R5,000'],
    ['Contingency', '5%', 'R13,120'],
]

PARTICIPANTS = [
    ['mizo_phyll', '2 dedicated shoot sessions, full identity on camera, closing performance, music licensing'],
    ['the_wife', 'Most intimate testimony, full identity, sensitive subject matter'],
    ['tman_gavin', 'Supporting witness, negotiated identity'],
    ['nicodemic', 'Supporting witness, negotiated identity'],
    ['the_producer', 'Supporting witness, full identity, origin story'],
    ['the_faith_witness', 'Single session, supporting role'],
]

SCHEDULE = [
    ['Day 1', 'Mizo Phyll — primary interview (60–90 min recorded)', 'Production studio, Thohoyandou'],
    ['Day 2', 'Wife, Tman Gavin, Nicodemic — individual interviews, isolated from each other', 'Production studio, Thohoyandou'],
    ['Day 3', 'DJ Davic + Faith witness — interviews', 'Production studio, Thohoyandou'],
    ['Day 3', 'Mizo Phyll — return interview (situational questions) + closing performance', 'Production studio, Thohoyandou'],
    ['Day 4', 'B-roll — Thohoyandou streets, Maniini Block J, VenRap environment', 'Thohoyandou locations'],
    ['Day 5', 'Dramatic reconstruction — early career scenes', 'Thohoyandou locations'],
    ['Day 6', 'Contingency / pickup shots', 'TBC'],
]

SUBJECTS = [
    ['Mizo Phyll (Livhuwani Aubrey Ratshiungo)', 'Principal subject — the spine', 'Thohoyandou', 'Full face'],
    ['Wife of Mizo Phyll', 'Most intimate voice', 'Thohoyandou', 'Full face'],
    ['Tman Gavin', 'Supporting witness', 'Thohoyandou', 'Negotiated'],
    ['Nicodemic', 'Supporting witness', 'Thohoyandou', 'Negotiated'],
    ['DJ Davic', 'Producer — origin story', 'Thohoyandou', 'Full face'],
    ['Faith community witness', 'Spiritual reckoning', 'Thohoyandou', 'Negotiated'],
]

CHECKLIST = [
    ['✓', 'Detailed treatment of proposed documentary (this document)'],
    ['✓', 'Description of locations and names of individuals to be interviewed (Sections 10 and 11)'],
    ['☐', 'Location release form — production studio, Thohoyandou'],
    ['☐', 'Location release forms — Thohoyandou b-roll locations'],
    ['☐', 'Subject release forms — all six principal contributors'],
    ['☐', 'Signed e.tv disclaimer form'],
    ['✓', 'Production budget with detailed personnel and resource breakdown (Section 9)'],
    ['✓', 'Confirmation of key personnel based in Limpopo province'],
    ['☐', 'Production company registration details'],
    ['☐', 'Names, race and gender of shareholders'],
    ['☐', 'Names, race and gender of directors'],
    ['☐', 'Names, race and gender of senior personnel and positions'],
    ['☐', 'Details of training initiatives in past 12 months'],
]


# ─── BUILD ───────────────────────────────────────────────────────────────────
def build():
    company = get_company()
    logline = get_logline()
    cq = get_cq()

    RIGHT = WD_ALIGN_PARAGRAPH.RIGHT

    # ── COVER ──
    sp(2400)
    centred('DOCUMENTARY PROPOSAL', 20, color=MID, bold=True)
    centred('e.tv — Hidden Gems of Mzansi: Regional Documentaries', 20, color=MID, after=80)
    sp(400); hr(ACCENT, 8); sp(200)
    centred('AGAINST ALL ODDS', 22, color=MID, bold=True)
    centred('VENDA NGA DECEMBER', 52, color=DARK, bold=True, before=120, after=120)
    centred('The King of VenRap', 22, color=MID, bold=True)
    sp(200); hr(ACCENT, 8); sp(400)
    centred('A Documentary Film', 20, color=MID)
    centred('Thohoyandou, Vhembe District, Limpopo', 20, color=MID, bold=True, before=60, after=60)
    sp(600)
    centred(company, 20, bold=True)
    centred('Limpopo, South Africa', 18, color=MID, before=80, after=80)
    sp(300)
    centred('Running time: 23 minutes  |  Language: Tshivenda / English — with English subtitles', 18, color=MID)
    centred('Proposed budget: R265,120.00', 18, color=MID, before=60, after=60)
    centred('Submission: documentaries@etv.co.za', 18, color=MID)
    page_break()

    # ── 1 LOGLINE ──
    h1('1. LOGLINE'); hr(); sp(80)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(10)
    _run(p, logline, size=12, color=DARK, bold=True, italic=True)
    sp(200)

    # ── 2 HIDDEN GEM ──
    h1('2. THE HIDDEN GEM — THOHOYANDOU AND VENRAP'); hr(); sp(80)
    body('South Africa knows Limpopo for the Kruger National Park, the baobab trees, and the Beit Bridge border. What it does not know — what it has never been shown — is that deep in the Vhembe District, in the streets and yards of Thohoyandou, a music culture grew without asking anyone\'s permission.')
    body('VenRap. Venda hip-hop. A sound that carries the Tshivenda language, the Limpopo landscape, the spiritual depth of a people, and the hunger of young men who grew up in the northernmost corner of the country and decided the world needed to hear them.')
    body('Most South Africans have never heard of VenRap. Most South Africans could not find Thohoyandou on a map. One man changed that. He carried Thohoyandou on his chest — literally — and took it to stages, studios, and collaborations that the north had never seen before.')
    body('This documentary is the story of that man, that place, and that journey. It is made by someone who was there at the beginning — who designed the first album cover, who watched from Musina before anyone believed. This access cannot be bought. No outsider can make this film.')
    sp(200)

    # ── 3 STORY ──
    h1('3. THE STORY'); hr(); sp(80)
    body('There is a tattoo on his chest. Three words: VENDA NGA DECEMBER.')
    body('He got it before the record deal. Before the national stages. Before the collaborations with artists the whole country knew. He got it when he was a young man from Thohoyandou with dreads, a deep faith, and a sound nobody had a name for yet. The tattoo was a declaration — a vow, a flag planted in his own skin before the world decided whether he was worth anything.')
    body('This documentary never leaves that tattoo. Every conversation, every testimony, every moment of the film radiates outward from those three words and returns to them. This is not a career documentary. It does not follow a timeline or trace a discography. It asks one question:')
    sp(80)
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Twips(720)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    _run(p, cq, size=11, color=DARK, italic=True)
    sp(120)
    body('The answer lives in his chest.')
    sp(200)

    # ── 4 SUBJECT ──
    h1('4. THE SUBJECT — MIZO PHYLL'); hr(); sp(80)
    body('Mizo Phyll — Livhuwani Aubrey Ratshiungo — is the King of VenRap. He says so himself when he introduces himself: Ndi dzi king. And he is not wrong.')
    body('He came from Maniini Block J, Thohoyandou, when VenRap was not yet a genre anyone recognised. He carried the Rastafari faith — the dreads, the doctrine, the vows. He made hip-hop that sounded like nowhere else in South Africa because it came from nowhere else.')
    body('Then the music industry found him. The industry had conditions. The dreads came off. What that meant — to him, to his faith community, to his wife, to the men who grew up alongside him in the same yards — is the emotional heart of this documentary.')
    body('He worked with the late Ricky Rick. He brought Venda to national stages. He created the Venda Nga December platform for Venda artists. He won the first-ever Best Tshivenda Hip-Hop category at the Tshivenda Music Awards in 2012. He has agreed in principle to participate — full face, full name, the tattoo on camera.')
    sp(200)

    # ── 5 CHARACTERS ──
    h1('5. CHARACTERS'); hr(); sp(80)
    char_slugs = [
        ['mizo_phyll', 'The Spine'],
        ['the_wife', 'The Most Intimate Voice'],
        ['tman_gavin', 'The Witness — Same Streets, Different Outcome'],
        ['nicodemic', 'The Witness — Same Streets, Different Outcome'],
        ['the_producer', 'The Producer Who Believed First'],
        ['the_faith_witness', 'The Spiritual Reckoning'],
        ['ricky_rick', 'The Absent Presence'],
    ]
    for slug, label in char_slugs:
        c = char_data(slug)
        if not c:
            continue
        h2('%s — %s' % (c['name'], label))
        for para in [x.strip() for x in re.split(r'\n+', c['role']) if x.strip()][:3]:
            body(para)
        if c['arch']:
            body('Verified facts for this character:', bold=True)
            for a in c['arch']:
                bul(a)
        sp(120)
    sp(200)

    # ── 6 ARC ──
    h1('6. NARRATIVE ARC — 23 MINUTES'); hr(); sp(80)
    for i in range(1, 8):
        s = scene_block(i)
        if not s:
            continue
        h3('%s%s' % (s['title'], ('  —  ' + s['runtime']) if s['runtime'] else ''))
        if s['desc']:
            body(s['desc'])
        sp(100)
    sp(200)

    # ── 7 VISUAL ──
    h1('7. VISUAL AND PRODUCTION APPROACH'); hr(); sp(80)
    h2('The Production Studio')
    body('All principal interviews are conducted in a single controlled studio environment in Thohoyandou — dressed and lit specifically for this documentary. No natural light. Complete control of the image. The studio creates visual unity across all characters, evokes the sacred consultation space of indigenous tradition, and eliminates location noise and continuity problems. This is a deliberate aesthetic decision, not a budget limitation.')
    h2('The Bubble Information Layer')
    body('This documentary introduces a second simultaneous information track — a bubble layer that surfaces verified facts, archived quotes, and community testimony throughout the film, without ever interrupting the speaker or cutting away from their face. Three types operate across the film:')
    bul('Archive Bubbles — verified historical facts appearing as clean text during informational moments. Example: "My African Dream — first artist to win Best Tshivenda Hip-Hop, Tshivenda Music Awards (2012)"')
    bul('Cut Floor Bubbles — extraordinary material from interviews that the 23-minute structure cannot hold, including the late Ricky Rick\'s own archived public words about Mizo Phyll. His own words, from the record. Not paraphrased. Not narrated.')
    bul('Community Bubbles — public social media testimony collected during pre-production. Facebook does not forget. These fill the dark spaces during the closing performance.')
    h2('Multi-Camera Setup')
    body('The studio operates on a multi-camera locked-off system. Multiple cameras are positioned and secured before each interview. The director functions as interviewer — fully present in the conversation, not managing equipment. Minimal crew of two in the room.')
    h2('The Chest — The Film\'s Recurring Image')
    body('The tattoo on his chest is the film\'s visual heartbeat. It is the opening image. It is the ad break transition device at every break. It is a detail shot during interview. It is the film\'s last image before black. VENDA NGA DECEMBER is the most repeated image in this documentary.')
    h2('The Closing Performance and Credits')
    body('The film ends with a live performance in the studio — alone, no audience, no stage, no direction. As he raps, community testimony fills the dark spaces around him. Then the messages transform — same visual format — into credits. The names of the people who made this film arrive in the same format as the facts that carried the whole film. The credits are the film\'s final argument: these names were also permanent.')
    h2('Dramatic Reconstruction')
    body('Brief reconstructions are used for moments that cannot be filmed — early career dismissal, the first performance nobody attended. Shot in a warmer, overexposed colour treatment. Silent, narrated. Local performers. One shoot day.')
    sp(200)

    # ── 8 PLAN ──
    h1('8. PRODUCTION PLAN'); hr(); sp(80)
    t = table([1200, 5760, 2400], 1 + len(SCHEDULE))
    h_row(t.rows[0], [('Day', 1200), ('Activity', 5760), ('Location', 2400)])
    for i, (d, a, l) in enumerate(SCHEDULE):
        d_row(t.rows[i + 1],
              [(d, 1200, {'bold': True}), (a, 5760),
               (l, 2400, {'italic': True, 'color': MID})], i % 2 != 0)
    sp(200)

    # ── 9 BUDGET ──
    h1('9. PRODUCTION BUDGET'); hr(); sp(80)
    t = table([3200, 4160, 2000], 2 + len(BUDGET))
    h_row(t.rows[0], [('Line Item', 3200), ('Detail', 4160), ('Amount', 2000)])
    for i, (item, detail, amount) in enumerate(BUDGET):
        d_row(t.rows[i + 1],
              [(item, 3200), (detail, 4160, {'italic': True, 'color': MID}),
               (amount, 2000, {'align': RIGHT})], i % 2 != 0)
    total = t.rows[len(BUDGET) + 1]
    cl(total.cells[0], 'TOTAL', 3200, fill=DARK, color=WHITE, bold=True, size=20)
    cl(total.cells[1], '', 4160, fill=DARK)
    cl(total.cells[2], 'R265,120', 2000, fill=DARK, color=WHITE, bold=True, size=20, align=RIGHT)
    sp(160)
    h2('Participant Fee Breakdown')
    t = table([3200, 4360, 1800], 1 + len(PARTICIPANTS))
    h_row(t.rows[0], [('Participant', 3200), ('Justification', 4360), ('Fee', 1800)])
    for i, (slug, just) in enumerate(PARTICIPANTS):
        c = char_data(slug)
        name = c['name'] if c else slug
        fee = c['fee'] if c else '—'
        d_row(t.rows[i + 1],
              [(name, 3200, {'bold': True}),
               (just, 4360, {'italic': True, 'color': MID}),
               (fee, 1800, {'align': RIGHT})], i % 2 != 0)
    sp(80)
    body('All participant fees: 50% on signed release form, 50% on completion of filming days. All payments via EFT with signed receipts.', italic=True, color=MID)
    sp(200)

    # ── 10 LOCATIONS ──
    h1('10. LOCATIONS'); hr(); sp(80)
    h2('Primary — Production Studio, Thohoyandou')
    body('Controlled studio environment — rented room in Thohoyandou, dressed and lit as a professional production studio. Used across Days 1, 2 and 3. Address confirmed on production company registration.')
    h2('Secondary — Thohoyandou Town and Surrounds')
    body('Maniini Block J and surrounding streets, Thohoyandou town centre, Vhembe District locations associated with the VenRap scene. All filming in public spaces or with confirmed location permission. Used for b-roll and dramatic reconstruction, Days 4 and 5.')
    sp(200)

    # ── 11 SUBJECTS ──
    h1('11. SUBJECTS TO BE INTERVIEWED'); hr(); sp(80)
    widths = [2800, 3160, 1800, 1600]
    t = table(widths, 1 + len(SUBJECTS))
    h_row(t.rows[0], [('Subject', 2800), ('Role', 3160), ('Location', 1800), ('Identity', 1600)])
    for i, r in enumerate(SUBJECTS):
        d_row(t.rows[i + 1], [(txt, widths[j]) for j, txt in enumerate(r)], i % 2 != 0)
    sp(80)
    body('Formal signed release forms to be submitted with the final proposal package. Mizo Phyll and DJ Davic have agreed in principle to participate. All other subjects have been approached.', italic=True, color=MID)
    sp(200)

    # ── 12 NO PRESENTER ──
    h1('12. PRESENTER / COMMENTATOR'); hr(); sp(80)
    body('This documentary does not use a presenter or narrator. The film is entirely built from the voices of its principal subjects and the bubble information layer. A presenter\'s voice would impose a singular interpretation on a story whose power comes from allowing multiple truths to coexist. The production studio is the presenter. The tattoo is the narrator. The audience is the judge.')
    sp(200)

    # ── 13 WHY ──
    h1('13. WHY THIS STORY. WHY THOHOYANDOU. WHY NOW.'); hr(); sp(80)
    body('South Africa has never seen a documentary about VenRap. It has never been shown Thohoyandou as a music capital. It does not know that a man from the streets of Vhembe carried his hometown tattooed on his chest all the way to national stages and never took it off.')
    body('This is the hidden gem. Not just the artist — the entire culture he came from and represents. A music tradition that grew without permission, without industry support, without anyone in Johannesburg paying attention — and produced a King.')
    body('Every South African who has ever come from somewhere the country overlooked will recognise this film. Every person who has ever had to negotiate between where they came from and where they are going will understand this film.')
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(10)
    _run(p, 'That is a national audience. And this is their hidden gem.', size=11, color=DARK, bold=True, italic=True)
    sp(200)

    # ── 14 COMPANY ──
    h1('14. PRODUCTION COMPANY INFORMATION'); hr(); sp(80)
    company_rows = [
        ['Company name', company],
        ['Province of operation', 'Limpopo'],
        ['Production base', 'Limpopo, South Africa'],
        ['Key personnel location', 'All key personnel are based in Limpopo province'],
        ['Registration number', '[To be completed]'],
        ['Director details', '[To be completed]'],
        ['Shareholder details', '[To be completed]'],
        ['BEE details', '[To be completed]'],
        ['Training initiatives (12 months)', '[To be completed]'],
    ]
    t = table([3600, 5760], len(company_rows))
    for i, (k, v) in enumerate(company_rows):
        d_row(t.rows[i], [(k, 3600, {'bold': True}), (v, 5760)], i % 2 != 0)
    sp(200)

    # ── 15 CHECKLIST ──
    h1('15. SUBMISSION DOCUMENTS CHECKLIST'); hr(); sp(80)
    t = table([720, 8640], len(CHECKLIST))
    for i, (check, text) in enumerate(CHECKLIST):
        row = t.rows[i]
        cl(row.cells[0], check, 720, no_border=True,
           color=(ACCENT if check == '✓' else DARK), bold=True,
           align=WD_ALIGN_PARAGRAPH.CENTER, pad_left=0, pad_right=120)
        cl(row.cells[1], text, 8640, no_border=True,
           fill=(WHITE if i % 2 == 0 else LIGHT))
    sp(200)

    # ── CLOSE ──
    hr(ACCENT, 6); sp(120)
    centred('AGAINST ALL ODDS  /  VENDA NGA DECEMBER  /  The King of VenRap', 20, bold=True)
    centred('A documentary film. Thohoyandou, Limpopo. 2026.', 18, color=MID, italic=True, before=80, after=80)
    centred(company, 18, bold=True)
    centred('documentaries@etv.co.za', 18, color=MID, before=60, after=60)


# ─── WRITE ───────────────────────────────────────────────────────────────────
def _rezip_deterministic(path):
    """Rewrite the docx zip deterministically so unchanged content produces
    byte-identical output.

    A .docx is a zip archive; python-docx stamps each entry with the current
    local time, so every run changes bytes even when content is identical —
    which makes CI's "commit only if changed" guard commit churn on every run.
    Re-writing the archive with sorted entries, a fixed DOS timestamp
    (1980-01-01, the zip epoch) and fixed compression makes the output a pure
    function of its content.
    """
    with zipfile.ZipFile(path) as zin:
        entries = [(name, zin.read(name)) for name in sorted(zin.namelist())]
    tmp = path + '.tmp'
    with zipfile.ZipFile(tmp, 'w', zipfile.ZIP_DEFLATED, compresslevel=9) as zout:
        for name, data in entries:
            info = zipfile.ZipInfo(name, date_time=(1980, 1, 1, 0, 0, 0))
            info.compress_type = zipfile.ZIP_DEFLATED
            info.external_attr = 0o600 << 16  # fixed -rw------- perms
            zout.writestr(info, data)
    os.replace(tmp, path)


def main():
    try:
        build()
        os.makedirs(os.path.dirname(OUT), exist_ok=True)
        doc.save(OUT)
        _rezip_deterministic(OUT)
    except Exception as err:  # noqa: BLE001
        sys.stderr.write('Build error: %s\n' % err)
        sys.exit(1)
    size = os.path.getsize(OUT)
    print('✅  Proposal written → %s' % OUT)
    print('    %.1f KB   |   Edit any file in film/%s/ and rerun.\n' % (size / 1024, project))


if __name__ == '__main__':
    main()
